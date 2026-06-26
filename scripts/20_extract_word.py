#!/usr/bin/env python3
import argparse
import json
import os
import sys
import traceback
from concurrent.futures import ProcessPoolExecutor, as_completed
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))

import mammoth
from docx import Document
from docx2python import docx2python
from tqdm import tqdm

from src.common import (
    CONVERTED_DIR,
    FAILED_DIR,
    OUTPUT_DIR,
    PARSED_WORD_DIR,
    LIBREOFFICE_CMD,
    WORD_EXTENSIONS,
    append_jsonl,
    clean_text,
    detect_language_from_text,
    ensure_base_dirs,
    get_quality_flags,
    iter_jsonl,
    run_cmd,
    write_json,
    write_jsonl,
)


REGISTRY_JSONL = OUTPUT_DIR / "document_registry.jsonl"

OUT_PARSED_JSONL = OUTPUT_DIR / "parsed_word.jsonl"
OUT_SUMMARY_JSON = OUTPUT_DIR / "parsed_word_summary.json"

FAILED_JSONL = FAILED_DIR / "word_failures.jsonl"


# ============================================================
# Native DOCX extractors
# ============================================================

def extract_docx_mammoth(path: Path) -> str:
    with path.open("rb") as f:
        result = mammoth.extract_raw_text(f)
    return clean_text(result.value)


def extract_docx_python_docx(path: Path) -> str:
    doc = Document(str(path))
    blocks = []

    for p in doc.paragraphs:
        txt = clean_text(p.text)
        if txt:
            blocks.append(txt)

    for table in doc.tables:
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                blocks.append(" | ".join(cells))

    return clean_text("\n".join(blocks))


def extract_docx_docx2python(path: Path) -> str:
    result = docx2python(str(path))
    try:
        text = result.text or ""
    finally:
        try:
            result.close()
        except Exception:
            pass

    return clean_text(text)


def extract_docx_native_best(path: Path) -> Tuple[str, str, Dict[str, int]]:
    candidates = []
    extractor_stats = {}

    for method_name, func in [
        ("docx_mammoth", extract_docx_mammoth),
        ("docx_python_docx", extract_docx_python_docx),
        ("docx_docx2python", extract_docx_docx2python),
    ]:
        try:
            text = func(path)
            text = clean_text(text)
            extractor_stats[method_name] = len(text)
            candidates.append((method_name, text))
        except Exception:
            extractor_stats[method_name] = -1

    if not candidates:
        raise RuntimeError(f"All native DOCX extractors failed: {path}")

    best_method, best_text = sorted(
        candidates,
        key=lambda x: len(x[1]),
        reverse=True,
    )[0]

    return best_text, best_method, extractor_stats


# ============================================================
# Docling DOCX extraction
# ============================================================

def extract_docx_docling(path: Path) -> Tuple[str, str, Dict[str, Any]]:
    """
    Returns:
      markdown_content, plain_text, docling_dict

    Docling is used here for structure preservation: headings, tables,
    lists, reading order. Native extractors are still kept as fallback.
    """

    from docling.datamodel.base_models import InputFormat
    from docling.document_converter import DocumentConverter, WordFormatOption
    from docling.pipeline.simple_pipeline import SimplePipeline

    converter = DocumentConverter(
        allowed_formats=[InputFormat.DOCX],
        format_options={
            InputFormat.DOCX: WordFormatOption(
                pipeline_cls=SimplePipeline,
            )
        },
    )

    result = converter.convert(str(path))
    doc = result.document

    markdown_content = doc.export_to_markdown()
    plain_text = doc.export_to_text()
    docling_dict = doc.export_to_dict()

    return (
        clean_text(markdown_content),
        clean_text(plain_text),
        docling_dict,
    )


# ============================================================
# DOC → DOCX conversion
# ============================================================

def libreoffice_convert_isolated(
    source_path: Path,
    target_ext: str,
    out_dir: Path,
    variant_id: str,
) -> Path:
    """
    Uses isolated LibreOffice profile to reduce parallel conversion conflicts.
    """

    out_dir.mkdir(parents=True, exist_ok=True)

    profile_dir = (CONVERTED_DIR / "lo_profiles" / variant_id).resolve()
    profile_dir.mkdir(parents=True, exist_ok=True)

    profile_uri = profile_dir.as_uri()

    cmd = [
        LIBREOFFICE_CMD,
        f"-env:UserInstallation={profile_uri}",
        "--headless",
        "--nologo",
        "--nofirststartwizard",
        "--convert-to",
        target_ext,
        "--outdir",
        str(out_dir),
        str(source_path),
    ]

    run_cmd(cmd, timeout=360)

    expected = out_dir / f"{source_path.stem}.{target_ext}"

    if expected.exists():
        return expected

    matches = list(out_dir.glob(f"{source_path.stem}*.{target_ext}"))

    if matches:
        return matches[0]

    raise FileNotFoundError(
        f"LibreOffice output not found: {source_path} → {target_ext}"
    )


def convert_doc_to_docx(source_path: Path, variant_id: str) -> Path:
    out_dir = CONVERTED_DIR / "doc_to_docx" / variant_id
    return libreoffice_convert_isolated(source_path, "docx", out_dir, variant_id)


# ============================================================
# Output helpers
# ============================================================

def make_markdown_fallback(title: str, text: str) -> str:
    text = clean_text(text)

    if not text:
        return f"# {title}\n"

    return f"# {title}\n\n{text}\n"


def write_word_outputs(
    variant_id: str,
    record: Dict[str, Any],
    markdown_content: str,
    docling_document: Optional[Dict[str, Any]],
) -> None:
    """
    Writes:
      data/parsed_word/{variant_id}.json
      data/parsed_word/{variant_id}.md
      data/parsed_word/docling_json/{variant_id}.docling.json  when available

    No TXT output is produced.
    """

    PARSED_WORD_DIR.mkdir(parents=True, exist_ok=True)

    md_path = PARSED_WORD_DIR / f"{variant_id}.md"
    json_path = PARSED_WORD_DIR / f"{variant_id}.json"

    md_path.write_text(markdown_content, encoding="utf-8")

    record["markdown_path"] = md_path.as_posix()

    if docling_document is not None:
        docling_json_dir = PARSED_WORD_DIR / "docling_json"
        docling_json_dir.mkdir(parents=True, exist_ok=True)

        docling_json_path = docling_json_dir / f"{variant_id}.docling.json"
        docling_json_path.write_text(
            json.dumps(docling_document, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        record["docling_json_path"] = docling_json_path.as_posix()
    else:
        record["docling_json_path"] = None

    write_json(json_path, record)


# ============================================================
# Worker
# ============================================================

def extract_one_word_variant(row: Dict[str, Any], min_chars: int) -> Dict[str, Any]:
    source_path = Path(row["source_path"])
    source_format = row["source_format"]
    variant_id = row["variant_id"]

    if not source_path.exists():
        raise FileNotFoundError(f"Source file not found: {source_path}")

    converted_path = None
    extraction_source_path = source_path

    if source_format == "doc":
        converted_path_obj = convert_doc_to_docx(source_path, variant_id)
        converted_path = converted_path_obj.as_posix()
        extraction_source_path = converted_path_obj

    elif source_format == "docx":
        extraction_source_path = source_path

    else:
        raise ValueError(f"Not a Word source_format: {source_format}")

    native_text = ""
    native_method = None
    native_stats: Dict[str, int] = {}

    docling_markdown = ""
    docling_text = ""
    docling_document = None
    docling_error = None

    # 1. Native Python extraction.
    try:
        native_text, native_method, native_stats = extract_docx_native_best(extraction_source_path)
    except Exception as e:
        native_method = "native_failed"
        native_stats = {}
        native_text = ""
        native_error = repr(e)
    else:
        native_error = None

    # 2. Docling extraction.
    # It is allowed to fail without killing the whole Word extraction,
    # because native extractors are the first safety layer.
    try:
        docling_markdown, docling_text, docling_document = extract_docx_docling(extraction_source_path)
    except Exception as e:
        docling_error = repr(e)
        docling_markdown = ""
        docling_text = ""
        docling_document = None

    candidates = []

    if native_text:
        candidates.append(
            {
                "method": native_method or "native_unknown",
                "source": "native",
                "text": native_text,
                "markdown": make_markdown_fallback(source_path.stem, native_text),
                "length": len(native_text),
            }
        )

    if docling_text:
        candidates.append(
            {
                "method": "docling_docx_simple_pipeline",
                "source": "docling",
                "text": docling_text,
                "markdown": docling_markdown or make_markdown_fallback(source_path.stem, docling_text),
                "length": len(docling_text),
            }
        )

    if not candidates:
        raise RuntimeError(
            f"Both native and Docling Word extraction failed: {source_path}. "
            f"native_error={native_error}, docling_error={docling_error}"
        )

    # Prefer longer extraction. In most clinical Word files, longer text means
    # fewer missing tables/sections.
    best = sorted(candidates, key=lambda x: x["length"], reverse=True)[0]

    raw_text = clean_text(best["text"])
    markdown_content = best["markdown"]

    detected_language = detect_language_from_text(raw_text)
    quality_flags = get_quality_flags(raw_text, detected_language)

    needs_review = False
    review_reason = None

    if len(raw_text) < min_chars:
        needs_review = True
        review_reason = "word_extraction_weak_low_char_count"

        if "low_char_count" not in quality_flags:
            quality_flags.append("low_char_count")

    if native_error:
        if "native_extraction_error" not in quality_flags:
            quality_flags.append("native_extraction_error")

    if docling_error:
        if "docling_extraction_error" not in quality_flags:
            quality_flags.append("docling_extraction_error")

    record = {
        **row,
        "detected_language": detected_language,
        "extraction_method": best["method"] if source_format == "docx" else f"doc_to_docx_{best['method']}",
        "extraction_source": best["source"],
        "converted_path": converted_path,
        "native_method": native_method,
        "native_char_count": len(native_text),
        "docling_char_count": len(docling_text),
        "extractor_stats": native_stats,
        "native_error": native_error,
        "docling_error": docling_error,
        "char_count": len(raw_text),
        "quality_flags": quality_flags,
        "needs_review": needs_review,
        "review_reason": review_reason,
        "needs_ocr": False,
        "ocr_reason": None,
        "raw_text": raw_text,
    }

    write_word_outputs(
        variant_id=variant_id,
        record=record,
        markdown_content=markdown_content,
        docling_document=docling_document,
    )

    return record


def safe_worker(args: Tuple[Dict[str, Any], int]) -> Tuple[str, Dict[str, Any]]:
    row, min_chars = args

    try:
        result = extract_one_word_variant(row, min_chars)
        return "ok", result

    except Exception as e:
        fail = {
            **row,
            "needs_review": True,
            "review_reason": "word_extraction_failed",
            "needs_ocr": False,
            "ocr_reason": None,
            "error": str(e),
            "traceback": traceback.format_exc(),
        }
        return "failed", fail


# ============================================================
# Load / rebuild
# ============================================================

def load_word_registry_rows(
    only_format: Optional[str] = None,
    only_unknown: bool = False,
) -> List[Dict[str, Any]]:
    rows = []

    for row in iter_jsonl(REGISTRY_JSONL):
        ext = row.get("source_extension")
        fmt = row.get("source_format")

        if ext not in WORD_EXTENSIONS:
            continue

        if only_format and fmt != only_format:
            continue

        if only_unknown and row.get("language_hint") != "unknown":
            continue

        rows.append(row)

    return rows


def read_json_safe(path: Path) -> Optional[Dict[str, Any]]:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return None


def rebuild_parsed_word_jsonl() -> int:
    rows = []

    for p in sorted(PARSED_WORD_DIR.glob("*.json")):
        row = read_json_safe(p)

        if row:
            rows.append(row)

    write_jsonl(OUT_PARSED_JSONL, rows)
    return len(rows)


# ============================================================
# Main
# ============================================================

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--limit", type=int, default=0, help="0 = no limit")
    parser.add_argument("--start", type=int, default=0)
    parser.add_argument("--workers", type=int, default=max(1, min(4, os.cpu_count() or 1)))
    parser.add_argument("--force", action="store_true")
    parser.add_argument("--min-chars", type=int, default=500)
    parser.add_argument("--only-format", choices=["docx", "doc"], default=None)
    parser.add_argument("--only-unknown", action="store_true")
    args = parser.parse_args()

    ensure_base_dirs()
    PARSED_WORD_DIR.mkdir(parents=True, exist_ok=True)

    rows = load_word_registry_rows(
        only_format=args.only_format,
        only_unknown=args.only_unknown,
    )

    rows = rows[args.start:]

    if args.limit and args.limit > 0:
        rows = rows[:args.limit]

    if not rows:
        print("No Word registry rows found.")
        return

    print(f"Word rows selected: {len(rows)}")
    print(f"Workers: {args.workers}")
    print(f"Min chars: {args.min_chars}")

    tasks = []
    skipped_existing = 0

    for row in rows:
        out_path = PARSED_WORD_DIR / f'{row["variant_id"]}.json'

        if out_path.exists() and not args.force:
            skipped_existing += 1
            continue

        tasks.append(row)

    parsed_ok = 0
    parsed_weak_needs_review = 0
    failed = 0

    if tasks:
        with ProcessPoolExecutor(max_workers=args.workers) as executor:
            futures = [
                executor.submit(safe_worker, (row, args.min_chars))
                for row in tasks
            ]

            for future in tqdm(
                as_completed(futures),
                total=len(futures),
                desc="Extracting Word",
            ):
                status, result = future.result()

                if status == "ok":
                    parsed_ok += 1

                    if result.get("needs_review"):
                        parsed_weak_needs_review += 1
                        append_jsonl(
                            FAILED_JSONL,
                            {
                                "variant_id": result["variant_id"],
                                "source_path": result["source_path"],
                                "source_format": result["source_format"],
                                "language_hint": result.get("language_hint"),
                                "detected_language": result.get("detected_language"),
                                "needs_review": True,
                                "review_reason": result.get("review_reason"),
                                "char_count": result.get("char_count"),
                                "quality_flags": result.get("quality_flags", []),
                            },
                        )

                else:
                    failed += 1
                    append_jsonl(FAILED_JSONL, result)

    total_parsed_files = rebuild_parsed_word_jsonl()

    summary = {
        "registry_word_rows_considered": len(rows),
        "tasks_run": len(tasks),
        "skipped_existing": skipped_existing,
        "parsed_ok": parsed_ok,
        "parsed_weak_needs_review": parsed_weak_needs_review,
        "failed": failed,
        "total_parsed_word_files": total_parsed_files,
        "outputs": {
            "parsed_word_jsonl": OUT_PARSED_JSONL.as_posix(),
            "parsed_word_dir": PARSED_WORD_DIR.as_posix(),
            "word_failures_jsonl": FAILED_JSONL.as_posix(),
            "summary_json": OUT_SUMMARY_JSON.as_posix(),
        },
    }

    write_json(OUT_SUMMARY_JSON, summary)

    print("\nDONE")

    for k, v in summary.items():
        if k != "outputs":
            print(f"{k}: {v}")

    print("\nSaved:")

    for _, v in summary["outputs"].items():
        print(v)


if __name__ == "__main__":
    main()
